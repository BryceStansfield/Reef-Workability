from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colours
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
WHITE_HEX      = "FFFFFF"
LIGHT_BLUE_HEX = "DEEAF1"
NAVY_HEX       = "1F4E79"
MID_BLUE_HEX   = "2E75B6"

# Column widths in DXA (1440 = 1 inch)
MODEL_W  = 2200
BRIER_W  = 1160
METRIC_W = 1120


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_borders(cell, color="AAAAAA"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _write_cell(cell, text, bold=False, font_color=WHITE, font_size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=WHITE_HEX, width_dxa=None):
    _set_cell_bg(cell, bg_hex)
    _set_cell_margins(cell)
    _set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if width_dxa:
        _set_col_width(cell, width_dxa)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = font_color
    run.font.name = "Arial"


def _merge_and_write(row, start_col, end_col, text, bg_hex, font_color=WHITE):
    cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[i])
    _write_cell(cell, text, bold=True, font_color=font_color,
                bg_hex=bg_hex, align=WD_ALIGN_PARAGRAPH.CENTER)


def generate_table(results: dict, output_path: str):
    """
    Generate a model performance comparison Word document.

    Args:
        results:     Dict keyed by model name. Each value must contain:
                         brier_skill_score, success_f1, success_precision,
                         success_recall, failure_f1, failure_precision, failure_recall.
                     Any extra keys (e.g. roc_curve, importances) are ignored.
        output_path: Path to save the .docx file, e.g. "results/comparison.docx"
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(11)
    section.page_height   = Inches(8.5)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Model Performance Comparison")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Arial"
    title_para.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=2 + len(results), cols=8)
    table.style = 'Table Grid'

    # Top-level headers
    r0 = table.rows[0]
    _write_cell(r0.cells[0], "Model",             bold=True, bg_hex=NAVY_HEX, width_dxa=MODEL_W)
    _write_cell(r0.cells[1], "Brier Skill Score", bold=True, bg_hex=NAVY_HEX, width_dxa=BRIER_W)
    _merge_and_write(r0, 2, 4, "Success", bg_hex=MID_BLUE_HEX)
    _merge_and_write(r0, 5, 7, "Failure", bg_hex=MID_BLUE_HEX)

    # Sub-headers
    r1 = table.rows[1]
    _write_cell(r1.cells[0], "", bold=True, bg_hex=NAVY_HEX, width_dxa=MODEL_W)
    _write_cell(r1.cells[1], "", bold=True, bg_hex=NAVY_HEX, width_dxa=BRIER_W)
    for i, label in enumerate(["F1", "Precision", "Recall", "F1", "Precision", "Recall"]):
        _write_cell(r1.cells[2 + i], label, bold=True, bg_hex=MID_BLUE_HEX, width_dxa=METRIC_W)

    # Data rows
    black = RGBColor(0, 0, 0)
    for row_idx, (model_name, metrics) in enumerate(results.items()):
        bg = WHITE_HEX if row_idx % 2 == 0 else LIGHT_BLUE_HEX
        row = table.rows[2 + row_idx]
        values = [
            model_name,
            f"{metrics['brier_skill_score']:.4f}",
            f"{metrics['success_f1']:.4f}",
            f"{metrics['success_precision']:.4f}",
            f"{metrics['success_recall']:.4f}",
            f"{metrics['failure_f1']:.4f}",
            f"{metrics['failure_precision']:.4f}",
            f"{metrics['failure_recall']:.4f}",
        ]
        widths = [MODEL_W, BRIER_W] + [METRIC_W] * 6
        aligns = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 7
        for col_idx, (val, w, al) in enumerate(zip(values, widths, aligns)):
            _write_cell(row.cells[col_idx], val, bold=False, font_color=black,
                        bg_hex=bg, align=al, width_dxa=w)

    doc.save(output_path)
    print(f"Saved to {output_path}")